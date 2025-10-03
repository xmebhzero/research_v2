import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

def create_styled_document(data):
    """
    Creates a Word document with specified styles and content.

    Args:
        data (dict): A dictionary containing the data to be inserted into the document.
                    Expected structure:
                    {
                        'project_name': str,
                        'transcriber': str,
                        'idi': str,
                        'city': str,
                        'execution_date': str,
                        'execution_time': str,
                        'moderator': str,
                        'respondent_name': str,
                        'dialogue': list of dicts  # See create_dialogue_section for details
                    }

    Returns:
        docx.document.Document: A Document object, ready to be saved.  Returns None on error.
    """
    if not isinstance(data, dict):
        print("Error: Input data must be a dictionary.")
        return None

    document = docx.Document()

    # 0. Document level Styles
    _set_document_margins(document, top=0.5, bottom=0.5, left=0.5, right=0.5)  # in inches

    # 1. Create metadata section
    _create_metadata_section(document, data)

    # 2. Create dialogue section
    _create_dialogue_section(document, data.get('dialogue', []))

    return document

def _set_document_margins(document, top, bottom, left, right):
    """Sets the document's margins.

    Args:
        document (docx.document.Document): The document object.
        top, bottom, left, right (float): Margin values in inches.
    """
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def _create_metadata_section(document, data):
    """Creates the metadata section of the document.

    Args:
        document (docx.document.Document): The document object.
        data (dict):  Dictionary containing the metadata.
    """
    # Create a table with 2 columns
    table = document.add_table(rows=0, cols=2)  # Start with 0 rows and 2 columns

    # Add data to the table
    metadata_pairs = [
        ("Nama Project", data.get('project_name', '')),
        ("Transcriber", data.get('transcriber', '')),
        ("IDI", data.get('idi', '')),
        ("Kota", data.get('city', '')),
        ("Tanggal pelaksanaan", data.get('execution_date', '')),
        ("Hari/Jam pelaksanaan", data.get('execution_time', '')),
        ("Moderator", data.get('moderator', '')),
        ("Nama Responden", data.get('respondent_name', ''))
    ]

    for label, value in metadata_pairs:
        row_cells = table.add_row().cells
        row_cells[0].text = f"{label}"
        row_cells[1].text = f"{value}"
        # Set font
        for cell in row_cells:
            cell.paragraphs[0].runs[0].font.name = 'Arial'

    # Add a space after the metadata table
    document.add_paragraph()
    

def _create_dialogue_section(document, dialogue_data):
    """Creates the dialogue section of the document.

    Args:
        document (docx.document.Document): The document object.
        dialogue_data (list of dict): A list where each dict represents a dialogue turn.
            Each dict should have:
            'speaker' (str):  'M' or 'R'
            'text'    (str): The dialogue text.
    """
    if not isinstance(dialogue_data, list):
        print("Error: Dialogue data must be a list.")
        return

    for turn in dialogue_data:
        if not isinstance(turn, dict) or 'speaker' not in turn or 'text' not in turn:
            print("Error: Each dialogue turn must be a dictionary with 'speaker' and 'text' keys.")
            continue

        speaker = turn.get('speaker', '')
        text = turn.get('text', '')

        p = document.add_paragraph()
        run = p.add_run(f"{speaker}\t\t") # Add two tabs
        run.font.name = 'Arial'
        run = p.add_run(f"{text}\n")
        run.font.name = 'Arial'
        p.paragraph_format.line_spacing = 1
        p.paragraph_format.space_after = Pt(0)

def main():
    """Main function to execute the document creation."""
    # Example data
    data = {
        'project_name': 'Acceptance',
        'transcriber': 'Mentari Meida',
        'idi': 'IDI 2',
        'city': 'Jakarta',
        'execution_date': '9 Oktober 2024',
        'execution_time': 'Rabu/10.00 WIB',
        'moderator': 'Janna',
        'respondent_name': 'Dedy',
        'dialogue': [
            {'speaker': 'M', 'text': 'oke, boleh nih Pak Dedy gantian kenalan, tinggalnya di mana dan kesibukannya ngapain sih pak?'},
            {'speaker': 'R', 'text': 'nama saya Dedy Firmansyah. Tinggal sih tergantung proyek ya, tapi rumah aslinya di Bekasi sebetulnya. Tapi ini sekarang lagi ada di Depok, soalnya deket sama gudang. Gudang kan lokasinya di Bojong Gede.'},
            {'speaker': 'M', 'text': 'oke oke.'},
            {'speaker': 'R', 'text': 'Alamat tinggal sih di Jati Kramat Bekasi'},
            {'speaker': 'M', 'text': 'hobinya apa nih pak?'},
            {'speaker': 'R', 'text': 'olahraga sama games sih paling'},
            {'speaker': 'M', 'text': 'senin sampe jumat kesibukannya ngapain?'},
            {'speaker': 'R', 'text': 'tergantung ya, kalau ada pengantaran, paling ya nyari-nyari schedule juga soalnya sekarang kebanyakan lagi pengantaran atau pemasangan ada di seminar-seminar itu senin sampe rabu itu lumayan sih. Kalau weekend, Jumat, Sabtu, Minggu, udah pasti.'},
            {'speaker': 'M', 'text': 'kerjanya jadi di weekend juga?'},
            {'speaker': 'R', 'text': 'untuk saat ini kebanyakan di weekend, saat ini. Ada di Departemen Agama sama Pariwisata biasanya sih di Senin sampe Rabu rata-rata'},
            {'speaker': 'M', 'text': 'oke oke'},
            {'speaker': 'R', 'text': 'soalnya gak bisa ditentuin secara rutin Senin-Kamis saya kerja, tapi kalau weekend pasti ada sih. Pemasangan-pemasangan multimedia untuk sekarang setelah pandemi lagi alhamdulillah banyak sih.'},
            {'speaker': 'M', 'text': 'kalau kantornya bapak seperti apa sih biasanya operasional sehari-harinya?'},
            {'speaker': 'R', 'text': 'sebenernya sih jatuhnya saya multimedia, tapi kebanyakan di luar sana sih bilangnya penyewaan videotron, kayak semacam LED, LCD, yang dipake buat dipanggung-panggung. Ada proyektor juga. Rata-rata yang kayak gitu. Proyektor buat seminar-seminar. Cuma emang pengennya sih megang panggung besar tapi alatnya belum ada semua. Dia kan harus memadai. Mentok-mentoknya sih semacam EO ya, cuma EO kan besar tuh, kalau saya kan sampe saat ini medianya masih videotron aja.'},
            {'speaker': 'M', 'text': 'itu biasanya dipake untuk acara-acara apa?'},
            {'speaker': 'R', 'text': 'seminar kebanyakan kalau sekarang. Kalau ibu lihat di panggung seminar kan ada yang pake LED, TV Digital, terus kursinya panjang kan suka ada monitor yang kanan kiri, nah itu biasanya ada di tempat saya gitu.'},
            {'speaker': 'M', 'text': 'jadi kliennya seperti apa?'},
            {'speaker': 'R', 'text': 'untuk saat ini vendor saya ada di Kementrian Agama sama Kementrian Pariwisata sih. Soalnya runag lingkup kita kan kecil ya. Kalau ada acara semua vendor jadi satu, ada bagian panggung, ada yang bagian kursi, gitu-gitu. Saya alhamdulillah masuk ke bagian multimedianya. Jadi setiap ada event di Departmen Agama atau Pariwisata alhamdulillah masih kepake.'},
            {'speaker': 'M', 'text': 'oke oke. Kalau sound sistem juga termasuk?'},
            {'speaker': 'R', 'text': 'termasuk, cuma alat saya masih terbatas. Jadi saya kerjasama sama temen gitu, partneran. Karena kalau kita cuma ngadain alat yang sebatas kita punya aja kan klien kadang maunya yang lengkap spek-speknya, sesuai yang mereka mau gitu. Jadi by request-nya mereka aja. Kalau barang saya masih mencukupi ya pake barang saya.'},
            {'speaker': 'M', 'text': 'posisi bapak sebagai apa sih?'},
            {'speaker': 'R', 'text': 'pemilik'},
            {'speaker': 'M', 'text': 'oh oke oke. Ownernya langsung ya.'},
            {'speaker': 'R', 'text': 'dibilang boss juga bukan sih, yang ngadain lah'},
            {'speaker': 'M', 'text': 'kepalanya ya?'},
            {'speaker': 'R', 'text': 'kalau saya boss udah punya gedung, ini mah masih rumahan'},
            {'speaker': 'M', 'text': 'gpp lah, itu kan juga usaha pak. Baik, kalau sehari-hari bapak sendiri atau ada tim nih?'},
            {'speaker': 'R', 'text': 'saya ada kru. Kalau sendiri ga bisa sih. Kalau ada proyek, udah sesuai jadwal, saya suruh kru saya. Pertama tentuin dulu, survey dulu tempatnya, kita lihat luasnya, terus kapasitas pandangannya. Media kan berhubungan sama pandangan penonton itu mau kemana arahnya. Kalau udah survey, udah sesuai sama yang kita hitung, berapa monitor yang kita butuhin, itu kan berhubungan juga sama pencahayaan nanti, lighting-nya, kalau mereka gak ada lighting itu kita tinggal bilang aja, “Mau gimana?”. Saya kan namanya bisnis ya, masukin barang saya juga gitu. Kalau saya udah survey, baru tentuin barang-barang yang dibutuhin sama sesuaiin dengan spek-spek yang mereka minta sih sebenernya.'},
            {'speaker': 'M', 'text': 'oke oke, flownya kayak gitu ya.'},
            {'speaker': 'R', 'text': 'kita kan gak bisa juga mereka rikues ibaratnya speknya gede tapi kapasitas ruangan gak mencukupi, kan sama aja, percuma juga, gak terpasang nanti. Mendingan kita kasih tau, luasnya segini cocoknya panggungnya sebesar apa, monitornya segimana.'},
            {'speaker': 'M', 'text': 'oke. Setelah kita udah tau nih tempatnya kayak gimana, udah ada denahnya, terus selanjutnya apa?'},
            {'speaker': 'R', 'text': 'pemasangan. Tapi pemasangan itu kan gak bisa di hari H. Misalnya acara tanggal 12 pagi. Kita jam 11 malam udah harus prepare, H-1. Itu bisa masuk barang kan tergantung tempatnya juga, bisa loading dock jam berapa. Rata-rata jam 10 malam udah bisa masuk. Nanti kita lihat set up mana yang... kalau acara kayak gitu kan biasanya panggung dulu ya, jadi bagian panggung dulu tuh yang ngerjain. Panggung selesai baru kita set up.'},
            {'speaker': 'M', 'text': 'bawa barang-barang itu dari gudang yang di Bojong Gede?'},
            {'speaker': 'R', 'text': 'iya,di Bojong Gede'},
            {'speaker': 'M', 'text': 'bawanya pake apa tuh pak barangnya?'},
            {'speaker': 'R', 'text': 'kalau saat ini paling pake, yang ada kan mobil saya van sama pick up ya. Barang-barang agak besar di mobil pick up, tapi barang-barang rentan kayak TV 60 inch, proyektor, itu saya masukin di van. Soalnya cuaca kan ngaruh juga. Pick up kalau hujan juga kan repot ya harus bongkar-bongkar pake terpal. Alhamdulillah saya punya van, jadi barang yang rentan saya masukin di van, jadinya kan ketutup.'},
            {'speaker': 'M', 'text': 'oke oke. Kalau di pick up barangnya seperti apa yang dibawa?'},
            {'speaker': 'R', 'text': 'stan-stannya monitor tuh, itu kan besi panjang. Terus box-box kalau misalnya ada tambahan sound sistem, itu kan pake box juga. Semacem itu lah.'},
            {'speaker': 'M', 'text': 'itu pake pick up? Kalau hujan gimana tuh?'},
            {'speaker': 'R', 'text': 'ya makanya pake terpal kan. Di jalan kan juga harus tertutup. Sama DISHUB kan pick up kalau di jalan itu juga harus tertutup, gak boleh terbuka.'},
            {'speaker': 'M', 'text': 'apa nih pengalaman yang kurang mengenakan selama bekerja di bidang ini gitu? Ada kejadian-kejadian gak?'},
            {'speaker': 'R', 'text': 'keadaan kendaraan sih sebenernya. Misalnya ada ban bocor, itu kan jadi memperlambat waktu ya. Walaupun acaranya besok tapi kan kita kadang udah ditungguin untuk pemasangan kan, harus tetep sesuai jadwal lah, kadang kan ada gladi resik. Ya itu sih kendalanya, di kendaraan sama cuaca sih kalau pick up. Walaupun pake terpal sebetulnya agak safety, tapi kalau hujannya deras banget, yang bawa itu kan suka ragu buat nerusin. Gak semua yang di pick up itu kan barang-barangnya ketutup semua, misalnya kayak kipas angin. Kalau cuaca lagi parah tapi ada barang yang terbuka, itu biasanya mereka agak takut buat nerusin. Kadang, “Gimana nih, pak?”. Jadi kalau waktunya masih mencukupi ya tunggu aja dulu reda. Tapi kalau seandainya waktunya gak mencukupi, ya disesuaikan sama jadwalnya juga sih. Mereka takutnya kan basah, terus koslet. Tapi kita kan sebetulnya bisa nentuin barang mana yang masuk ke pick up. Cuma mereka kadang suka agak takut.'},
            {'speaker': 'M', 'text': 'oke baik. Kalau kejadian yang menguntungkan, yang menyenangkan, bisnis multimedia saat ini tuh seperti apa sih? Cukup oke kah atau gimana nih?'},
            {'speaker': 'R', 'text': 'untuk saat ini, pasca pandemi, alhamdulillah cukup banyak ya. Dulu kan saya ikut orang ya, itu sebelum pandemi. Nah, ketika saya ikut orang kan saya ngambil ilmunya juga, dan butuh modal juga sih emang. Awalnya proyek mereka tuh lumayan, hampir full. Ibaratnya sebulan itu bisa di angka 15 hari tuh kegiatannya mereka. 15 hari tuh full semua acara, di luar weekend ya. Itu yang jadi saya mikir, oh enak banget, proyeknya banyak gitu. Walaupun dulu kita masih bawahan, capek, tapi kita bisa ngelihat lah prospek ke depannya gimana. Karena di multimedia ini gak semena-mena cuma di seminar atau pertemuan, tapi ternyata bisa dipake ke musik. Sebelum pandemi ya, dulu kan panggung ya cuma panggung tok. Kalau sekarang semua panggung ada LED semua di belakangnya, terus permainan cahaya. Dari situ saya lihat awalnya. Yaudah sebelum pandemi saya buka sendiri. Saya ngumpulin dulu tuh baru saya keluar dari mereka. Dan kalau di bidang kita ya, sebenernya ruang lingkupnya kecil, karena setiap event timnya itu itu aja sebenernya. Kita nyebutnya “Lo lagi, Lo lagi”. Dari situ awalnya saya ijin dulu saya mau buka. Katanya ya udah gpp. Karena kan ruang lingkupnya kecil, kalau saya buka, mereka pasti tau juga. Udah berkembang, mereka udah pesat lah, udah ada PT yang naungin, saya mainnya di wilayah yang pernah saya kerjasama sama mereka. Alhamdulillah saya dikasih pegang Departement Agama sama Pariwisata. Alhamdulillah udah ngumpulin barang-barang semua, udah ada kerjasama sama vendor mereka, jadi setiap ada acara saya yang dipake di situ. Awalnya ngerintis sih sodara yang saya bawa buat bantuin sama temen. Sekarang kalau masang kayak begitu berdua doang udah gak bisa, harus ada krunya, banyak.'},
            {'speaker': 'M', 'text': 'oke oke oke. Nah, selanjutnya aku mau tanya soal topik utamanya ya, kendaraan niaga. Tadi kan udah bilang ya, Pak Dedy tuh pakenya pick up sama van. Nah, waktu memutuskan untuk beli mobil ini tuh pertimbangannya apa aja sih?'},
            {'speaker': 'R', 'text': 'pertimbangan awalnya ya budget. Waktu itu awalnya saya kenapa ngambil van sama pick up karena yang saya lihat semuanya kan pake mobil tertutup. Waktu itu saya di awal-awal ada budget cuma untuk Grandmax van, yang tahun 2018. Ketika barang tinggi kayak stan-stan itu, besi-besi tinggi, itu kan gak masuk, makanya saya beli mobil pick up itu. Cuma memang awalnya sih pengennya sih kalau ada ya harusnya sih mobil engkel. Itu kan bisa masuk semua, gak perlu dua mobil. Tapi kan ya budget juga, harus memikirkan pengeluaran di tahap awal itu kan.'},
            {'speaker': 'M', 'text': 'oke, jadi yang pertama budget harus sesuai dulu nih ya sama kantong kita. Tapi selain itu ada pertimbangan lain lagi gak nih?'},
            {'speaker': 'R', 'text': 'kenapa saya pilih Grandmax waktu itu karena... Grandmax ini sebenernya waktu awal saya masih kerja sama mereka, saya buka sendiri ya. Terus saya kepikiran, mereka tuh sebetulnya udah ada mobil yang besar.'},
            {'speaker': 'M', 'text': 'sorry, mereka tuh siapa ya?'},
            {'speaker': 'R', 'text': 'tempat dulu saya kerja itu, yang sebelum saya misah dari mereka.'},
            {'speaker': 'M', 'text': 'oke oke'},
            {'speaker': 'R', 'text': 'mobil saya ini bisa saya pake buat supply dari gang sempit ke mobil mereka. Sebetulnya itu awalnya kenapa dipake. Saya langsung milih Grandmax karena fleksbile, bisa masuk ke gang-gang kecil. Kan banyak tuh acara di... kadang di daerah kota aja, di Slipi gitu kan mobil-mobil besar kadang gak bisa masuk kan. Nah, kalau pake Grandmax tuh masih bisa masuk. Awalnya sih gitu. Cuma memang kapasitas di dalam mobilnya itu gak bisa banyak karena terbatas sama ruang atap di atasnya juga kan.'},
            {'speaker': 'M', 'text': 'tadi kan bapak bilang di grandmax isinya kayak TV, terus apa lagi tuh?'},
            {'speaker': 'R', 'text': 'proyektor, box-box buat kabel, barang-barang rentan saya masukin ke van'},
            {'speaker': 'M', 'text': 'itu kan sebetulnya bisa pake mobil biasa ya, mobil penumpang, tapi kenapa milihnya pake Grandmax?'},
            {'speaker': 'R', 'text': 'kalau mobil penumpang, misalnya ibu bawa TV, itu kan harus pake dus ya. TV kan minimal di 30 inch, sekarang rata-rata pake 50-60 inch, itu kan agak tinggi. Kalau pake minibus biasa kayak Avanza, itu bisa masuk, tapi kan ada kursi, jadi kalau dipaksa ya bisa masuk tapi kapasitasnya tetep gak akan bisa banyak. Dan itu kalau misalnya peraturan kita bawa barang tapi bawa mobil yang kaca, yang keliatan dari luar lah, tapi itu full barang, itu sebetulnya gak boleh bu kalau di peraturan jalan.'},
            {'speaker': 'M', 'text': 'oh itu ada ya aturannya begitu?'},
            {'speaker': 'R', 'text': 'ada, ada aturannya. Ibu bawa Avanza nih, kardusnya sampe atas dan kaca spion tengah itu gak bisa ngeliat ke belakang, itu sebetulnya bisa ditangkap polisi. Ada aturannya itu.'},
            {'speaker': 'M', 'text': 'oh oke oke.'},
            {'speaker': 'R', 'text': 'makanya kenapa saya ambil mobil barang, jadi walaupun kita bawa barang sebanyak apa pun di dalam, selama gak melebihi kapasitasnya, aman. Pernah juga bawa Avanza sama van, Avanza tuh full, tapi kacanya kan paling 60, itu keliatan, diberhentiin sama polisi, kena tilang. Alasannya kan bukan mobil barang tapi kenapa dipenuhin sama  barang. TV itu kan kardusnya tinggi ya, itu ketutupan kan spionnya, jadi gak boleh.'},
            {'speaker': 'M', 'text': 'oke, jadi tadi karena bisa fleksible nyelip sama kapasitasnya juga memang mobil barang gitu ya?'},
            {'speaker': 'R', 'text': 'iya, betul.'},
            {'speaker': 'M', 'text': 'oke oke'},
            {'speaker': 'R', 'text': 'kalau mau fleksible sih ya mobil biasa emang enak buat dipake juga. Tapi kan saya jatuhnya buat pengiriman juga, itu kan ada aturannya juga kalau di jalan untuk mobil-mobil pengiriman.'},
            {'speaker': 'M', 'text': 'oke, terus apa lagi nih pertimbangan bapak pilih Grandmax?'},
            {'speaker': 'R', 'text': 'ya lebih simple sih'},
            {'speaker': 'M', 'text': 'bahan bakar gimana?'},
            {'speaker': 'R', 'text': 'bahan bakar menurut saya agak lebih irit kalau gak pake AC'},
            {'speaker': 'M', 'text': 'tapi biasanya pake AC atau enggak?'},
            {'speaker': 'R', 'text': 'tergantung moment aja sih. Anak-anak kan perokok juga, setau saya gak pake AC. Tapi pake AC kalau pas hujan sih.'},
            {'speaker': 'M', 'text': 'tapi itu gpp tuh? Kan barang-barang yang dibawa cukup sensitif ya, ada TV, ada proyektor, kalau kena suhu panas gpp?'},
            {'speaker': 'R', 'text': 'tergantung pengiriman. Kalau loading barang siang hari, itu pasti pake AC. Kan jaga kualitas barang juga. Van kan juga panas ya. Tapi kan selama ini jadwal pemasangan kita selalu malam.'},
            {'speaker': 'M', 'text': 'oke oke oke.'},
            {'speaker': 'R', 'text': 'semua mobil sih sama aja sebetulnya menurut saya ya. Tapi kalau yang saya alamin antara pick up sama van yang saya punya ya lebih safety di van. Buat bahan bakar juga lebih irit.'},
            {'speaker': 'M', 'text': 'kenapa van dianggap lebih safe?'},
            {'speaker': 'R', 'text': 'barang kan terlindungi kalau pake van, tapi kapasitas pick up lebih luas karena gak ada batasan kan kalau ke atas. Kalau van lebih safety aja. Kalau mau juga ada mobil engkel box, itu lebih enak menurut saya.'},
            {'speaker': 'M', 'text': 'kalau pick up, merknya apa pak?'},
            {'speaker': 'R', 'text': 'Grandmax Pick Up. Kan ada'},
            {'speaker': 'M', 'text': 'oh oke oke, sama ya brandnya. Kirain beda. Pertimbangannya apa tuh?'},
            {'speaker': 'R', 'text': 'ya itu aja, kapasitas. Kalau bawa tiang kan gak bisa masuk di van, soalnya ada yang dua meter. Itu sih bisa dilepas pasang, cuma kan kadang kita butuh buru-buru, makan waktu juga kan kalau lepas pasang lagi. Bracket-bracket buat di dinding, itu kan lumayan, kalau ditaro di pick up juga aman, besi semua kan. Cuaca basah juga gak masalah.'},
            {'speaker': 'M', 'text': 'kan pick up lebih bawa barang-brang besar dan berat ya, ngaruh juga gak sih ke pertimbangan kita nyari pick up yang tenaganya itu kuat gitu atau remnya lebih pakem?'},
            {'speaker': 'R', 'text': 'pasti, kalau mobil pasti kayak gitu. Kalau menurut saya setiap orang beli pasti dicek semuanya ya, mesin, rem, semua harus bagus ya. Cuma kalau ngeliat lebih pake Grandmax atau mobil lain, menurut saya sih itu tergantung pemakaian mobil ya sama service rutinnya.'},
            {'speaker': 'M', 'text': 'kalau kayak after salesnya dipikirin juga gak, pak?'},
            {'speaker': 'R', 'text': 'ketika beli sih enggak ya, tapi ketika mau upgrade ya pasti ada. Saya pasti harus ngalahin salah satu, kecuali budgetnya cukup tanpa harus menjual lagi.'},
            {'speaker': 'M', 'text': 'kalau pick up kan merknya ada banyak ya. Tapi kenapa waktu itu milihnya Grandmax?'},
            {'speaker': 'R', 'text': 'identik ya awalnya. Di tahun-tahun 2018 yang ada sih kebanyakan itu. Dulu sih ada Suzuki Carry, tapi itu kan kecil ya pick up nya. Kalau sekarang sih agak lebih besar ya. Terus juga ada Isuzu, itu ada. Mitsubishi juga, sebetulnya sih dulu pengen Mitsubishi yang Colt Diesel itu. Kenapa awalnya ke Grandmax karena itu masih bisa saya saving cost lah dari modal awal. Daripada kita beli yang lebih besar, awalnya kan saya ngelihat untuk merambah ke dunia ini kan saya masih nebak-nebak lah ibaratnya, bakal lanjtu atau enggak, gitu kan.'},
            {'speaker': 'M', 'text': 'jadi itu tergolongnya udah agak mahal ya, pak?'},
            {'speaker': 'R', 'text': 'bukan mahal sih, maksudnya waktu awal yang kita lihat kan mobil yang dipake itu itu aja, yang banyak di jalan lah ibaratnya. Kenapa gak pake merk lain, ya waktu itu saya pengennya van. Van dulu kan cuma ada setau saya Suzuki sama Daihatsu, MPV sama Grandmax ini. Akhir-akhir ini kan baru ada Wuling tuh ngeluarin van juga. Kalau saya kenapa pilih Grandmax, Suzuki MPV... karakter mobilnya sih ya, saya lebih suka di Grandmax, keliatannya lebih luas gitu.'},
            {'speaker': 'M', 'text': 'MPV gimana?'},
            {'speaker': 'R', 'text': 'lebih kecil dan kapasitasnya juga gak begitu luas'},
            {'speaker': 'M', 'text': 'oke oke oke. Jadi van sama pick up ini belinya tahun berapa?'},
            {'speaker': 'R', 'text': 'antara 2018-2019. Beli van dulu baru pick up.'},
            {'speaker': 'M', 'text': 'oke.'},{'speaker': 'R', 'text': 'Itu sekitar selish setahun dua tahunan lah sampe akhirnya beli pick up'},
            {'speaker': 'M', 'text': 'itu dua-duanya tapi belinya sebelum pandemi ya?'},
            {'speaker': 'R', 'text': 'iya'},
            {'speaker': 'M', 'text': 'dulu pas pak dedy milih-milih pick up ini, pertama kali sumber informasinya tuh dari mana sih untuk dapetin informasi yang lebih detail?'},
            {'speaker': 'R', 'text': 'dari rekanan dulu'},
            {'speaker': 'M', 'text': 'rekan kerja yang dulunya satu kantor?'},
            {'speaker': 'R', 'text': 'bukan, bukan. Gimana ya? Di kerjaan saya kan semua megang vendor masing-masing, termasuk saya. Vendor-vendor ini kan tadi saya bilang orangnya itu itu aja. Jadi saya lihat mereka bawa barang apa, pake mobil apa. Cuma memang kalau sekarang yang mobilnya paling kecil ya masih saya.'},
            {'speaker': 'M', 'text': 'oke baik. Jadi tau pertama kali dari rekanan ya. Yang ditanya dulu apa tuh pak?'},
            {'speaker': 'R', 'text': 'kebanyakan rekomen mereka sih.'},
            {'speaker': 'M', 'text': 'mereka ngerekomendasiin Grandmax itu ya?'},
            {'speaker': 'R', 'text': 'enggak, mereka gak ngerekomen jenis apa, merk apa, sebetulnya enggak. Cuma karena tipenya yang kita cari van, di tahun itu kan yang saya tau cuma ada Daihatsu sama MPV itu, Suzuki, yang berbentuk van. Tapi saya lebih tertarik ke Daihatsu itu.'},
            {'speaker': 'M', 'text': 'selain dari temen rekanan itu, cari infonya dari mana lagi?'},
            {'speaker': 'R', 'text': 'cari sendiri. Jaman itu kan medsos gak seperti sekarang ini ya, jadi teteplah showroom kita datengin. Ya paling mulut ke mulut lah.'},
            {'speaker': 'M', 'text': 'Oke. Nah, aku mau lebih detail lagi ya. Van kan pake Grandmax ya. Selama pake van ini gimana pengalamannya?'},
            {'speaker': 'R', 'text': 'untuk pemakaian? Biasa aja sih. Tapi kalau mobil, memang saya ada service rutin. Ini kan mobil gerak terus, tiap hari mereka keluar. Kalau service gak rutin kan nanti susah juga. Mobil biasanya kalau udah mesin satu kena, kan merembet lah ibaratnya, sama kayak motor matic.'},
            {'speaker': 'M', 'text': 'servicenya per bulan, per minggu atau gimana tuh?'},
            {'speaker': 'R', 'text': 'per kilometer saya sekarang. Awalnya di 5000 km, tapi kan konsultasi ke mekanik bengkel juga, jadi di angka 10.000 baru saya service'},
            {'speaker': 'M', 'text': 'kalau service di mana? Bengkel Daihatsu langsung atau di mana?'},
            {'speaker': 'R', 'text': 'awalnya di Daihatsu, tapi sekarang kan kita pengen yang lebih fleksible, ada langganan lah ibaratnya yang udah biasa megang tuh mobil.'},
            {'speaker': 'M', 'text': 'fleksible dalam segi apa nih pak?'},
            {'speaker': 'R', 'text': 'dari segi waktu. Kalau bengkel resmi kan ada yang gak 24 jam ya. Walaupun sekarang sih setau saya udah ada CS nya, tinggal telpon. Tapi kita ngelihat waktu dan efisien, bisa anter jemput, ya udah kita pake langganan ini.'},
            {'speaker': 'M', 'text': 'yang pake van ini bapak sendiri atau ada yang lain?'},
            {'speaker': 'R', 'text': 'tergantung. Kalau saya lagi mau ikut ke sana juga, saya di van. Saya lebih nyaman di van daripada di pick up.'},
            {'speaker': 'M', 'text': 'kenapa tuh pak?'},
            {'speaker': 'R', 'text': 'ya ada history nya juga kan. Itu kan yang pertama kali saya beli, yang pertama kali saya bawa waktu itu.'},
            {'speaker': 'M', 'text': 'oke oke'},
            {'speaker': 'R', 'text': 'jadi kalau mereka mau masang, saya ketemu EO nya, saya ikut mereka. Tapi kadang saya nyusul aja.'},
            {'speaker': 'M', 'text': 'bapak gak harus ada di tempat ya?'},
            {'speaker': 'R', 'text': 'gak harus, kan udah tau kerjaannya masing-masing. Paling tinggal kontrol di akhir doang.'},
            {'speaker': 'M', 'text': 'terus gimana nih pengalamannya bawa van Daihatsu?'},
            {'speaker': 'R', 'text': 'cukup sih untuk saat ini, paling kapasitas aja yang masih terbatas kalau van. Ruangannya ya. Kenapa saya pengen upgrade yang lebih gede ya karena itu, biar bisa masuk semua.'},
            {'speaker': 'M', 'text': 'upgrade ke van yang lebih gede atau ke engkel box?'},
            {'speaker': 'R', 'text': 'pengennya sih ke engkel box'},
            {'speaker': 'M', 'text': 'tapi kalau pake engkel box berarti van ini dijual atau gimana tuh?'},
            {'speaker': 'R', 'text': 'pengennya sih tetap, nambah lah gitu. Tapi belum tau ke depannya gimana. Pengennya sih nambah.'},
            {'speaker': 'M', 'text': 'tapi selama ini kapasitas van yang terbatas itu masih oke atau sampe harus bikin kita bolak balik ke venue?'},
            {'speaker': 'R', 'text': 'kalau barangnya yang kita bawa agak banyak ya pasti bolak balik. Tapi kita duluin barang yang mau dipasang duluan kan. Kalau ada barang gak muat, yang pokoknya dulu harus kepasang baru selanjutnay dikirim lagi.'},
            {'speaker': 'M', 'text': 'dari segi keamanannya, van daihatsu ini gimana?'},
            {'speaker': 'R', 'text': 'ya aman.'},
            {'speaker': 'M', 'text': 'pernah ada keluhan gak dari temen-temen yang bawa?'},
            {'speaker': 'R', 'text': 'belum ada keluhan. Van ini kan masih belum power stering, masih manual... kalau anak-anak sekarang kan bawanya yang udah power stering, untuk yang di setir ya, kalau van saya punya itu kan belum power stering, masih manual, jadi agak berat. Karena kru saya banyakan anak tahun 2000an mereka bilangnya berat bawa itu.'},
            {'speaker': 'M', 'text': 'oke itu ya keluhannya. Terus apa lagi nih?'},
            {'speaker': 'R', 'text': 'kalau nyaman saya lebih nyaman pake van. Dari semua keadaan di dalam mobilnya, saya lebih nyaman di van. Kalau pic up kan belakangnya tertutup, ngeliatnya ke arah luar kan. Kalau van kan nge-los, jadi keliatan barang masih ada apa enggak.'},
            {'speaker': 'M', 'text': 'oke oke oke. Terus apa lagi nih keluhannya soal van ini?'},
            {'speaker': 'R', 'text': 'ya paling apa ya? Kalau di mobil karena ya kita lumayan rutin service, paling ngeluhnya kalau pas pecah ban aja sih'},
            {'speaker': 'M', 'text': 'tapi sering tuh ban bocor?'},
            {'speaker': 'R', 'text': 'jarang sih, gak terlalu sering kalau ban, cuma kadang kalau udah aus mereka suka ngeluh, “Nih ban udah botak, gak diganti?” karena kan mereka juga harus jaga safety di jalan. Kalau ngeluh mah lebih ke uang jalan aja mereka. Kalau kendaraan kan kita emang harus safety, rutin sevice aja sih.'},
            {'speaker': 'M', 'text': 'tapi yang selama bapak lihat mereka nyaman gak bawa van ini?'},
            {'speaker': 'R', 'text': 'kalau belum terbiasa emang kurang begitu suka. Itu aja sih sebenernya. Tapi kalau udah biasa, malah lebih nyaman bawa Grandmax. Mereka juga ngakuin, gak ngerti juga kenapa. Mungkin karena mobilnya yang ada juga itu aja ya, daripada bawa pick up mereka lebih suka bawa Grandmax. Awalnya emang pada gak suka, berat. Tapi mereka di jalan kan lebih banyak ngelawan cuaca, jadi capek bawa pick up, harus pasang terpal dulu, gitu gitu sih.'},
            {'speaker': 'M', 'text': 'oke oke. Kalau pick up nih, keluhannya apa lagi nih selain tadi kita gak bisa lihat ke belakang?'},
            {'speaker': 'R', 'text': 'paling kendalanya di cuaca, barang-barang juga kadang... lampu merah aja sering ada orang-orang kayak pengamen gitu, mereka jadi curiga, itu kan gampang diambil juga kan. Walaupun besi tebal gitu, kalau kita lengah dikit, kalau kata anak-anak jalanan mah, dibetak, ya cepet juga.'},
            {'speaker': 'M', 'text': 'tapi pernah ada kejadian kayak gitu?'},
            {'speaker': 'R', 'text': 'enggak sih alhamdulillah. Tapi waspada aja sih, kita kan jalan selalu malam. Jalan kan berdua juga, jadi lebih waspada aja.'},
            {'speaker': 'M', 'text': 'yang kita lakuin apa nih biar barang-barang di pick up itu lebih aman?'},
            {'speaker': 'R', 'text': 'paling kita tutup terpal. Jadi kan barang lebih terlindungi lah, gak keliatan dari luar juga kan wlaupun secara bentuk kelihatan. Cuma kalau kita udah tau daerahnya, tanpa terpal sih bisa. Kita kan harus tau wilayahnya juga sih, sama cuaca juga.'},
            {'speaker': 'M', 'text': 'terus ada lagi gak keluhannya kalau di pick up?'},
            {'speaker': 'R', 'text': 'karena pick up ini kan agak baru, udah power stering juga, gak masalah sih. Paling ban bocor aja, nelpon gitu.'},
            {'speaker': 'M', 'text': 'tapi kalau dari pak dedy, ngelihat sekilas aja kru di lapangan, rata-rata lebih suka pake yang mana?'},
            {'speaker': 'R', 'text': 'van. Karena lebih fleksible kan, mereka masukin barang, selesai. Gak perlu ikat-ikat, gak perlu pasang terpal. Kalau pick up sebetulnya lebih enak kalau bawa barang besar, panjang, gitu. Itu lebih efisien pake pick up. Memang itu muatan pick up lah.'},
            {'speaker': 'M', 'text': 'pick up bapak sekarang bahan bakarnya gimana? irit gak?'},
            {'speaker': 'R', 'text': 'kalau irit atau enggak sebetulnya biasa aja saya lihat. Saya kan gak tau juga sebenernya, kan yang bawa mereka, yang pake mereka, nanti reimburse. Awalnya saya emang ada catatan, tapi ke sini sini saya udah tau lah berapa KM tuh butuh berapa banyak bensin. Jadi normal aja sih. Kalau irit apa enggaknya saya juga gak tau kan perbandingannya sama yang lain gimana. Yang saya pake kan dua ini aja.'},
            {'speaker': 'M', 'text': 'oke oke oke. Kalau pick up ini pake AC juga?'},
            {'speaker': 'R', 'text': 'ada, ada AC. 2018 ke sini udah pada AC semua.'},
            {'speaker': 'M', 'text': 'kalau segi kenyamanannya ada masalah gak?'},
            {'speaker': 'R', 'text': 'enggak sih.'},
            {'speaker': 'M', 'text': 'oke baik. Nah, kalau kita ngomongin pick up sama van, merk-merk apa aja sih yang bapak tau?'},
            {'speaker': 'R', 'text': 'saat ini ya? Dulu jaman saya masih dikit ya, Carry, Colt Diesel. Sekarang kan banyak. Isuzu Traga tuh, itu saya baru ngeliatnya. Kan Traga ada tuh yang semacam mobil box, dia engkel sama van, dia ada tuh. Saya lihat sih baru, baru di tahun 2020-an ini ya. Kalau di pasar kan banyak Mitsubishi colt diesel tapi itu terlalu besar kapasitasnya. Waktu itu kan gudang saya di Bekasi masih sempit, gak masuk itu. Kalau sekarang sih mau gak mau ya harus upgrade mobil sih.'},
            {'speaker': 'M', 'text': 'oke, terus merknya apa lagi nih?'},
            {'speaker': 'R', 'text': 'Suzuki ya paling. Wuling juga ngeluarin van kan sekarang. Saya pernah tau ada yang bawa. Bentuknya van. Dia depannya kayak Kijang Innova, tapi belakangnya van, itu juga ngeluarin.'},
            {'speaker': 'M', 'text': 'dia pick up?'},
            {'speaker': 'R', 'text': 'bukan, van'},
            {'speaker': 'M', 'text': 'liat di mana tuh pak?'},
            {'speaker': 'R', 'text': 'di jalan. Terus pas nganter, temen ada pake aplikasi, itu ada Wuling Van. Pas mereka order pake aplikasi waktu itu.'},
            {'speaker': 'M', 'text': 'Wuling Van untuk angkut barang juga?'},
            {'speaker': 'R', 'text': 'setau saya pastinya iya. Yang berhubungan sama blind van biasanya buat barang.'},
            {'speaker': 'M', 'text': 'oke oke'},
            {'speaker': 'R', 'text': 'mobil mobil yang ada KIR nya itu udah pasti mobil barang lah ibaratnya.'},
            {'speaker': 'M', 'text': 'ah oke oke oke. Nah, ini menurut bapak aja ya, sepengetahuan bapak, Mitsubishi itu brand yang seperti apa sih?'},
            {'speaker': 'R', 'text': 'kalau saya, Mitsubishi lebih wah. Kalau mobil Mitsubishi, dulu kan saya pengen Lancer, itu kan harganya agak tinggi. Tapi kalau di niaga, saya belum tau nih, belum bisa mengetahui secara dalam maksudnya. Mitsubishi kan ada mobil engkel yang gede juga kan. Tapi saya belum tau detail lah poin-poinnya untuk yang niaga ya.'},
            {'speaker': 'M', 'text': 'tapi temen bapak ada yang pake Mitsubishi juga untuk anter barang?'},
            {'speaker': 'R', 'text': 'ada yang pake engkel. Rata-rata yang besar itu'},
            {'speaker': 'M', 'text': 'kalau Isuzu, brand yang seperti apa nih?'},
            {'speaker': 'R', 'text': 'Isuzu juga ada yang pake. Nah itu yang saya bilang, waktu itu saya pernah liat temen saya pake Traga, itu saya tertarik.'},
            {'speaker': 'M', 'text': 'kenapa tertariknya?'},
            {'speaker': 'R', 'text': 'kapasitasnya dia juga lebih besar kan.'},
            {'speaker': 'M', 'text': 'oh oke oke. Terus apa lagi nih yang membuat kita tertarik ke Traga?'},
            {'speaker': 'R', 'text': 'tampilan sih. Yang Grandmax ini kan udah mobil jadul, 2018. Kalau mobil baru kan tampilannya lebih bagus. Sebetulnya saya sih gak terlalu lihat merk ya, saya pengen upgrade ya biar kapasitas lebih besar aja maksudnya.'},
            {'speaker': 'M', 'text': 'poin pentingnya itu ya?'},
            {'speaker': 'R', 'text': 'iya'},
            {'speaker': 'M', 'text': 'kalau tampilan Isuzu lebih modern ya berarti?'},
            {'speaker': 'R', 'text': 'iya untuk saat ini. Dia kan baru ya, baru tahun dua ribuan dia ngeluarin Traga itu. Kan ada juga tuh Isuzu yang semacem engkel tapi lebih kecil, di atas mobil box, tapi di bawah engkel. Susah saya jelasinnya tuh. Pokoknya lebih kecil dari engkel. Waktu itu saya tertarik ke situ tuh.'},
            {'speaker': 'M', 'text': 'oke oke oke. Kalau Suzuki gimana nih?'},
            {'speaker': 'R', 'text': 'paling yang saya tau yang pick up ya sekarang.'},
            {'speaker': 'M', 'text': 'yang Carry?'},
            {'speaker': 'R', 'text': 'iya. Kan sekarang lebih besar dibandingin Carry yang dulu, dia ngeluarin. Dan pick up nya lebih besar dibandingkan Grandmax kapasitas baknya itu lho. Tapi detailnya saya belum tau semana.'},
            {'speaker': 'M', 'text': 'antara merk-merk ini semua, ada van nya semua?'},
            {'speaker': 'R', 'text': 'Daihatsu ada, Isuzu ada. Mitsubishi setau saya mobil box ya. Gak tau juga dia ada van apa enggak, mungkin saya yang kurang ngeh, tapi setau saya mobil box aja. Kalau pengguna Mitsubishi kan setau saya yang colt diesel itu kan ya, baknya juga ada, terus mobil box juga ada. Kebanyakan mereka kan buat antar kota. Di dalam kota jarang.'},
            {'speaker': 'M', 'text': 'Isuzu ada gak van nya?'},
            {'speaker': 'R', 'text': 'kayanya belum ada ya, adanya box setau saya. Panther box kan ada. Kalau van saya belum tau.'},
            {'speaker': 'M', 'text': 'oke baik baik. Habis ini saya mau share screen. Nah, Isuzu Traga PU. Ini yang bapak liat di jalan bukan?'},
            {'speaker': 'R', 'text': 'belakangnya ada boksnya. Ini kan pick up ya. Saya lihat yang ada boxnya.'},
            {'speaker': 'M', 'text': 'yang cocok buat kebutuhan bisnis bapak yang mana?'},
            {'speaker': 'R', 'text': 'tergantung. Kenapa saya pilih van sama box kan tergantung kapasitas. Kalau ada box, lebih cocok sih ini.'},
            {'speaker': 'M', 'text': 'cocok ini? kalau ada boxnya?'},
            {'speaker': 'R', 'text': 'makanya tergantung... saya liat kapasitas sih. Panjang barang kan gak semua box bisa sampe, ada batas volume'},
            {'speaker': 'M', 'text': 'nah ini speknya pak. Gimana nih? Bisa dilihat aja dulu'},
            {'speaker': 'R', 'text': 'oh ini masuk semua kalau barang saya. Panjang, lebar, tinggi, masuk semua ini. Ini dimensi pick up belakangnya kan? Bukan panjang mobilnya kan?'},
            {'speaker': 'M', 'text': 'maunya panjang boxnya aja ya?'},
            {'speaker': 'R', 'text': 'karena kita berhubungan sama barang, kita lihat dimensi baknya ya biasanya.'},
            {'speaker': 'M', 'text': 'Oke oke. Kalau ini dimensi kapasitasnya, ini masuk gak sama barang-barang kita?'},
            {'speaker': 'R', 'text': 'masuk. Kapasitas muatan juga masuk, itu kan hampir 1,5 ton ya.'},
            {'speaker': 'M', 'text': 'kalau kecepatannya?'},
            {'speaker': 'R', 'text': 'untuk ukuran itu masuk. Mungkin kalau barang kecepatan itu kan maksimal 80 juga udah kenceng ya. Masuk ini, termasuk kenceng lah ibaratnya.'},
            {'speaker': 'M', 'text': 'oke oke. Nah kalau dengan spek begini, bapak anggap mobil Traga PU ini berapa harganya?'},
            {'speaker': 'R', 'text': '300-an juta.'},
            {'speaker': 'M', 'text': 'kenapa 300 juta?'},
            {'speaker': 'R', 'text': 'ya karena... ini gambaran saya aja sih'},
            {'speaker': 'M', 'text': 'iya gpp'},
            {'speaker': 'R', 'text': 'karena dia kan agak besar ya, terus dia bensin atau diesel sih?'},
            {'speaker': 'M', 'text': 'gak ada detailnya. Maunya apa?'},
            {'speaker': 'R', 'text': 'sama aja sih bensin apa diesel. Kalau mobil besar kan lebih irit diesel ya.'},
            {'speaker': 'M', 'text': 'oke oke. Jadi 300 juta itu diesel dan dimensinya seperti ini ya. Kalau mahal, semahal apa sih yang sampe bapak gak mau beli?'},
            {'speaker': 'R', 'text': 'tergantung speknya. Saya kan belum tau detail banget ya. Kalau saya lihat ini sih 300 juta masuk. Semua mobil kan beda-beda ya, unggulnya juga beda-beda. Kalau unggul ya kenapa enggak, kalau saya.'},
            {'speaker': 'M', 'text': 'kalau Daihatsu pick up bapak sekarang harganya juga di rentang 300 juta?'},
            {'speaker': 'R', 'text': 'oh enggak, masih di bawah. Grandmax aja saya beli 85 juta.'},
            {'speaker': 'M', 'text': 'oke baik. Nah sekarang Isuzu NMR L. Masih cocok gak buat pak dedy?'},
            {'speaker': 'R', 'text': 'untuk kapasitas saya... sebetulnya semakin besar makin cocok'},
            {'speaker': 'M', 'text': 'tapi ada limitnya juga kan pasti'},
            {'speaker': 'R', 'text': 'untuk saat ini, ini terlalu besar dimensinya untuk usaha saya. Kalau nanti saya bisa ngerambah ke EO, ini cocok, karena semua barang bisa masuk.'},
            {'speaker': 'M', 'text': 'kalau ini, dimensinya gak cocoknya di mananya tuh?'},
            {'speaker': 'R', 'text': 'kepanjanganya ya kalau buat saya'},
            {'speaker': 'M', 'text': 'harganya berapa nih kira-kira?'},
            {'speaker': 'R', 'text': 'kalau misalnya... tar dulu... kalau mobil engkel gitu kan di angka 400an lah ya.'},
            {'speaker': 'M', 'text': 'maksimalnya berapa nih yang kita ga mau beli?'},
            {'speaker': 'R', 'text': 'di angka 500, itu udah pasti. Kisaran saya sih di 400. Cuma ya itu, saya belum tau detailnya mobil-mobil gede. Belum menelusuri lah.'},
            {'speaker': 'M', 'text': 'oke gpp gpp. Selanjutnya ada Isuzu NMR HD 5.8.'},
            {'speaker': 'R', 'text': 'engkel ya?'},
            {'speaker': 'M', 'text': 'iya, tapi dimensinya lebih kecil dikit dari yang tadi'},
            {'speaker': 'R', 'text': 'setau saya ini buat barang-barang berat kayak bebatuan, pasir, gitu.'},
            {'speaker': 'M', 'text': 'buat barang bapak?'},
            {'speaker': 'R', 'text': 'cocoknya yang tadi'},
            {'speaker': 'M', 'text': 'yang NMR ya?'},
            {'speaker': 'R', 'text': 'iya, lebih safety'},
            {'speaker': 'M', 'text': 'karena ada boxnya?'},
            {'speaker': 'R', 'text': 'iya. Kalau mobil ini kan setau saya yang berhubungan sama medan ya, batuan, tanah. Ini kan lihat dari bentuk baknya ya, kalau agak tebal biasanya buat barang barang berat.'},
            {'speaker': 'M', 'text': 'kalau harganya kira-kira berapa nih?'},
            {'speaker': 'R', 'text': 'sama kayak tadi'},
            {'speaker': 'M', 'text': '400an?'},
            {'speaker': 'R', 'text': 'iya'},
            {'speaker': 'M', 'text': 'kenapa 400an?'},
            {'speaker': 'R', 'text': 'yang sebelumnya, yang 300an'},
            {'speaker': 'M', 'text': 'traga?'},
            {'speaker': 'R', 'text': 'iya, lebih mahal yang tadi kalau saya lihat'},
            {'speaker': 'M', 'text': 'traga kan 300an tadi katanya. Yang 5.8 berapa tuh?'},
            {'speaker': 'R', 'text': '400an. 350 lah, di bawah engkel tadi. Tapi speknya lebih... karena bahan. Bahannya dia lebih tebal'},
            {'speaker': 'M', 'text': 'yang 5.8?'},
            {'speaker': 'R', 'text': 'iya, dibanding yang pick up pertama ya?'},
            {'speaker': 'M', 'text': 'tapi kenapa harganya mirip yang NMR L ini pak?'},
            {'speaker': 'R', 'text': 'di atas yang L sih'},
            {'speaker': 'M', 'text': 'jadi berapa tuh?'},
            {'speaker': 'R', 'text': '350-450 lah'},
            {'speaker': 'M', 'text': 'oke oke oke. Sekarang yang 6.5. Beda di kecepatannya ya, dia lebih lambat. Kira-kira harganya gimana?'},
            {'speaker': 'R', 'text': 'lebih mahal sih'},
            {'speaker': 'M', 'text': 'kenapa lebih mahal?'},
            {'speaker': 'R', 'text': 'saya lihat dari bentuknya sama bagian belakangnya aja kan udah ada tambahan tuh kan, itu setau saya juga agak lebih tebal'},
            {'speaker': 'M', 'text': 'cocoknya buat angkut apa nih?'},
            {'speaker': 'R', 'text': 'batu, pasir, bahan bahan bangunan sih kalau saya lihat'},
            {'speaker': 'M', 'text': 'harganya gimana?'},
            {'speaker': 'R', 'text': 'ya pasti lebih mahal daripada yang biasa ya'},
            {'speaker': 'M', 'text': 'tadi kan yang 5.8 katanya maksimalnya 450 ya. Kalau yang 6.5 berapa nih?'},
            {'speaker': 'R', 'text': '400-500 sih'},
            {'speaker': 'M', 'text': '600 kemahalan gak?'},
            {'speaker': 'R', 'text': 'ya bisa jadi sampe 600. Kan kapasitasnya lebih besar juga'},
            {'speaker': 'M', 'text': 'kapasitasnya beda 6 kg sih lebih kecil dikit'},
            {'speaker': 'R', 'text': 'ya gak jauh beda lah, masih di angka 400an'},
            {'speaker': 'M', 'text': 'oke oke. Terakhir NMR FVM UHP. Cocoknya bawa barang apa nih?'},
            {'speaker': 'R', 'text': 'barang besar ya, besi atau yang berhubungan sama...'},
            {'speaker': 'M', 'text': 'cocok gak buat barang-barang bapak?'},
            {'speaker': 'R', 'text': 'gak cocok kalau saya'},
            {'speaker': 'M', 'text': 'kenapa tuh?'},
            {'speaker': 'R', 'text': 'besar. Barang saya juga bukan buat mobil kayak gini.'},
            {'speaker': 'M', 'text': 'dimensinya ini lebih wah lagi ya'},
            {'speaker': 'R', 'text': 'besar lah kalau ini'},
            {'speaker': 'M', 'text': 'harganya gimana menurut bapak kalau yang ini?'},
            {'speaker': 'R', 'text': 'bisa di angka 700-800 ke atas'},
            {'speaker': 'M', 'text': 'oke I see. Nah, kalau misalkan pak dedy, tadi kan di awal sempet bilang mau upgrade mobil. Itu di kondisi bisnis seperti apa sih sehingga kita memutuskan mau upgdare kendaraan?'},
            {'speaker': 'R', 'text': 'ya penambahan kerjaan sih. Maksudnya lagi padat-padatnya kerjaaan. Kadang kan mobil lagi keluar, tapi saya ambil kerjaan lain juga di waktu yang bersamaan. Beda proyek di tempat yang berbeda gitu.'},
            {'speaker': 'M', 'text': 'Oke oke oke. Terus apa lagi nih?'},
            {'speaker': 'R', 'text': 'ya paling kapasitas sih. Kadang di suatu event dia minta... kadang saya bawa 6 TV, waktu itu di JCC tuh, terus bawa bracket-bracket, terus projector, itu kan gak muat. Jadi pernah ada suatu moment semua barang di gudang itu keluar semua, itu ga muat di dua mobil.'},
            {'speaker': 'M', 'text': 'kalau kayak gitu gimana tuh pak?'},
            {'speaker': 'R', 'text': 'saya suruh yang urgent dibawa duluan. Bolak balik. Kalau bener-bener urgent waktunya, ya paling kita cari mobil lain. Sewa atau online gitu.'},
            {'speaker': 'M', 'text': 'kalau beli mobil baru nih, promo seperti apa yang menarik dan cocok buat kita?'},
            {'speaker': 'R', 'text': 'harga, diskon itu pasti.'},
            {'speaker': 'M', 'text': 'diskon seperti apa maunya?'},
            {'speaker': 'R', 'text': 'ya pasti kan diskon itu kan tergantung dari harga mobil ya dan kita tau harga awal mobilnya berapa. Harga 300 dapat 20% aja kan udah lumayan.'},
            {'speaker': 'M', 'text': 'kalau bapak maunya cash atau nyicil sih?'},
            {'speaker': 'R', 'text': 'untuk saat ini cash. Kemungkinan bisa cicilan. Insya allah kan tahun depan, tergantung budget juga lah nanti. Untuk saat ini dilakuin penyicilan sih gak masalah cuma saya masih ngelihat proyek saya di tahun depan gimana. Kalau urgent butuh upgrade ya pasti nyicil. Daripada saya sewa-sewa mobil lain, mendingan saya nyicil kan tapi mobil kan jadi milik.'},
            {'speaker': 'M', 'text': 'kalau bapak misalkan bisa membuatt mobil khusus bisnisnya bapak, kira-kira harus ada fitur apa aja tuh di mobil tersebut?'},
            {'speaker': 'R', 'text': 'fitur di dalamnya ya? AC.'},
            {'speaker': 'M', 'text': 'AC sentral atau gimana nih?'},
            {'speaker': 'R', 'text': 'AC sentral tergantung dari mobil juga ya saya liat. Kalau yang mobil barang kan lebih enak AC di depan karena orangnya hanya ada di depan. Tapi kalau ada sentralnya ya lebih bagus. Terus ada TV nya. Dulu cuma radio, sekarang udah ada TV. Kamera mundur juga. Dulu kan van saya belum ada, harus pasang manual. Terus kalau ada kamera mundur kan lebih efisien juga. Sekarang  kan maunya yang simpel simpel, ada semua. Mobil sekarang kan ada semua. Dulu kan mobil barang ya cuma mobil bawa barang, gak ada fitur fitur macem-macem, kamera mundur, AC, TV gitu gitu.'},
            {'speaker': 'M', 'text': 'jadi harapannya mobil barang juga manusiawi ya?'},
            {'speaker': 'R', 'text': 'iya, saya lihat udah ada sih kayak gitu. Gak tau bawaan dari pabriknya atau pasang sendiri. Saya kenapa tertarik ke Traga karena dari bentukannya terus sam fitur saya lihat juga udah lumayan ya.'},
            {'speaker': 'M', 'text': 'lumayan gimana tuh?'},
            {'speaker': 'R', 'text': 'kabinnya tuh saya lihat lebih enak lah'},
            {'speaker': 'M', 'text': 'enak gimana?'},
            {'speaker': 'R', 'text': 'pas saya masuk mau cobain Traga itu, enak juga nih mobil'},
            {'speaker': 'M', 'text': 'traga ada TV nya?'},
            {'speaker': 'R', 'text': 'kemarin sih saya lihat ada, tapi saya gak tanya itu bawaan atau rakit sendiri. Saya gak nanya ke situ sih.'},
            {'speaker': 'M', 'text': 'oke kalau gitu, itu pertanyaan saya terakhir.'}
        ]
    }

    # Create the document
    document = create_styled_document(data)

    if document:
        # Save the document
        document.save("transcript.docx")
        print("Document 'transcript.docx' created successfully.")
    else:
        print("Failed to create the document.")

if __name__ == "__main__":
    main()
